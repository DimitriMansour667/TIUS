import customtkinter as ctk
from tkinter import filedialog, messagebox, simpledialog
import tkinter as tk
import os
import sys
from pathlib import Path
from datetime import datetime
import json
import threading
from PIL import Image, ImageTk
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import zipfile
import shutil

# Configuration de CustomTkinter
ctk.set_appearance_mode("dark")  # Mode sombre par défaut
ctk.set_default_color_theme("blue")  # Thème bleu

class DocumentBinder:
    def __init__(self, root):
        self.root = root
        self.root.title("Relieur Virtuel de Documents")
        
        # Configuration en plein écran
        self.root.state('zoomed')  # Plein écran sous Windows
        # Alternative pour d'autres systèmes
        # self.root.attributes('-fullscreen', True)
        
        # Bind pour sortir du plein écran avec Escape
        self.root.bind('<Escape>', self.toggle_fullscreen)
        self.root.bind('<F11>', self.toggle_fullscreen)
        
        # Configuration du thème
        self.setup_theme()
        
        # Variables
        self.documents = []
        self.chapters = []
        self.output_path = ""
        self.current_doc_index = 0
        self.generation_cancelled = False  # Pour le contrôle de la génération PDF
        
        # Configuration
        self.config = {
            "default_output_dir": os.path.expanduser("~/Documents/Relieur_Virtuel"),
            "supported_formats": {
                "pdf": ["pdf"],
                "word": ["docx", "doc"],
                "image": ["png", "jpg", "jpeg", "bmp", "gif", "tiff"]
            },
            "page_size": "A4",
            "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1}
        }
        
        self.create_widgets()
        self.create_menu()
        
        # Créer le dossier de sortie par défaut
        os.makedirs(self.config["default_output_dir"], exist_ok=True)
    
    def setup_theme(self):
        """Configuration du thème CustomTkinter"""
        # Variables de thème
        self.colors = {
            "primary": "#1f538d",
            "secondary": "#14375e",
            "accent": "#3b82f6",
            "success": "#10b981",
            "warning": "#f59e0b",
            "error": "#ef4444",
            "text": "#ffffff",
            "text_secondary": "#9ca3af",
            "background": "#1a1a1a",
            "surface": "#2d2d2d",
            "border": "#404040"
        }
        
        # Fonts
        self.fonts = {
            "title": ("Segoe UI", 20, "bold"),
            "heading": ("Segoe UI", 16, "bold"),
            "body": ("Segoe UI", 12),
            "caption": ("Segoe UI", 10)
        }
    
    def create_menu(self):
        """Créer la barre de menu"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # Menu Fichier
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Fichier", menu=file_menu)
        file_menu.add_command(label="Nouveau projet", command=self.new_project)
        file_menu.add_command(label="Ouvrir projet", command=self.open_project)
        file_menu.add_command(label="Sauvegarder projet", command=self.save_project)
        file_menu.add_separator()
        file_menu.add_command(label="Quitter", command=self.root.quit)
        
        # Menu Edition
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Edition", menu=edit_menu)
        edit_menu.add_command(label="Paramètres", command=self.show_settings)
        
        # Menu Aide
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Aide", menu=help_menu)
        help_menu.add_command(label="Guide d'utilisation", command=self.show_help)
        help_menu.add_command(label="À propos", command=self.show_about)
    
    def create_widgets(self):
        """Créer l'interface utilisateur moderne avec CustomTkinter"""
        # Frame principal
        main_frame = ctk.CTkFrame(self.root, corner_radius=0)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Header avec titre et boutons de thème
        header_frame = ctk.CTkFrame(main_frame, height=80)
        header_frame.pack(fill="x", padx=10, pady=(10, 0))
        header_frame.pack_propagate(False)
        
        # Titre principal
        title_label = ctk.CTkLabel(
            header_frame, 
            text="📚 Relieur Virtuel de Documents",
            font=self.fonts["title"],
            text_color=self.colors["accent"]
        )
        title_label.pack(side="left", padx=20, pady=20)
        
        # Boutons de thème
        theme_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        theme_frame.pack(side="right", padx=20, pady=20)
        
        # Bouton thème sombre/clair
        self.theme_button = ctk.CTkButton(
            theme_frame,
            text="🌙 Sombre",
            width=100,
            height=35,
            command=self.toggle_theme,
            fg_color=self.colors["secondary"],
            hover_color=self.colors["primary"]
        )
        self.theme_button.pack(side="right", padx=5)
        
        # Barre d'actions
        action_frame = ctk.CTkFrame(main_frame)
        action_frame.pack(fill="x", padx=10, pady=10)
        
        # Boutons d'action avec icônes
        self.action_buttons = {}
        
        actions = [
            ("📄 Ajouter PDF", self.add_pdf, self.colors["primary"]),
            ("📝 Ajouter Word", self.add_word, self.colors["accent"]),
            ("🖼️ Ajouter Images", self.add_images, self.colors["success"]),
            ("📋 Organiser Chapitres", self.organize_chapters, self.colors["warning"]),
            ("🔨 Générer PDF", self.generate_pdf, self.colors["error"])
        ]
        
        for i, (text, command, color) in enumerate(actions):
            btn = ctk.CTkButton(
                action_frame,
                text=text,
                command=command,
                width=150,
                height=40,
                font=self.fonts["body"],
                fg_color=color,
                hover_color=self.adjust_color(color, -20)
            )
            btn.pack(side="left", padx=5, pady=10)
            self.action_buttons[text] = btn
        
        # Frame principal de contenu
        content_frame = ctk.CTkFrame(main_frame)
        content_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        
        # Frame gauche - Liste des documents
        left_frame = ctk.CTkFrame(content_frame)
        left_frame.pack(side="left", fill="both", expand=True, padx=(10, 5), pady=10)
        
        # Titre de la section
        docs_title = ctk.CTkLabel(
            left_frame,
            text="📋 Documents ajoutés",
            font=self.fonts["heading"],
            anchor="w"
        )
        docs_title.pack(fill="x", padx=15, pady=(15, 10))
        
        # Frame pour la liste des documents
        docs_list_frame = ctk.CTkFrame(left_frame, fg_color="transparent")
        docs_list_frame.pack(fill="both", expand=True, padx=15, pady=(0, 15))
        
        # Scrollable frame pour les documents
        self.docs_scrollable = ctk.CTkScrollableFrame(
            docs_list_frame,
            width=400,
            height=400,
            corner_radius=10
        )
        self.docs_scrollable.pack(fill="both", expand=True)
        
        # Boutons de contrôle des documents
        docs_controls = ctk.CTkFrame(left_frame)
        docs_controls.pack(fill="x", padx=15, pady=(0, 15))
        
        control_buttons = [
            ("⬆️ Monter", self.move_up, self.colors["accent"]),
            ("⬇️ Descendre", self.move_down, self.colors["accent"]),
            ("🗑️ Supprimer", self.remove_document, self.colors["error"])
        ]
        
        for text, command, color in control_buttons:
            btn = ctk.CTkButton(
                docs_controls,
                text=text,
                command=command,
                width=100,
                height=35,
                font=self.fonts["caption"],
                fg_color=color,
                hover_color=self.adjust_color(color, -20)
            )
            btn.pack(side="left", padx=5, pady=5)
        
        # Frame droit - Aperçu
        right_frame = ctk.CTkFrame(content_frame)
        right_frame.pack(side="right", fill="both", expand=True, padx=(5, 10), pady=10)
        
        # Titre de l'aperçu
        preview_title = ctk.CTkLabel(
            right_frame,
            text="👁️ Aperçu du document",
            font=self.fonts["heading"],
            anchor="w"
        )
        preview_title.pack(fill="x", padx=15, pady=(15, 10))
        
        # Frame pour l'aperçu
        preview_container = ctk.CTkFrame(right_frame, fg_color="transparent")
        preview_container.pack(fill="both", expand=True, padx=15, pady=(0, 15))
        
        # Canvas pour l'aperçu avec scrollbars
        self.preview_canvas = tk.Canvas(
            preview_container,
            bg="#2d2d2d",
            highlightthickness=0,
            relief="flat"
        )
        self.preview_canvas.pack(fill="both", expand=True)
        
        # Scrollbars pour l'aperçu
        preview_v_scrollbar = ctk.CTkScrollbar(
            preview_container,
            orientation="vertical",
            command=self.preview_canvas.yview
        )
        preview_v_scrollbar.pack(side="right", fill="y")
        self.preview_canvas.configure(yscrollcommand=preview_v_scrollbar.set)
        
        # Barre de statut
        self.status_frame = ctk.CTkFrame(main_frame, height=40)
        self.status_frame.pack(fill="x", padx=10, pady=(0, 10))
        self.status_frame.pack_propagate(False)
        
        self.status_var = tk.StringVar()
        self.status_var.set("✅ Prêt - Ajoutez des documents pour commencer")
        
        self.status_label = ctk.CTkLabel(
            self.status_frame,
            textvariable=self.status_var,
            font=self.fonts["caption"],
            anchor="w"
        )
        self.status_label.pack(fill="x", padx=20, pady=10)
        
        # Liste pour stocker les widgets de documents
        self.document_widgets = []
    
    def adjust_color(self, color, adjustment):
        """Ajuster la luminosité d'une couleur"""
        try:
            # Convertir la couleur hex en RGB
            color = color.lstrip('#')
            r, g, b = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
            
            # Ajuster la luminosité
            r = max(0, min(255, r + adjustment))
            g = max(0, min(255, g + adjustment))
            b = max(0, min(255, b + adjustment))
            
            return f"#{r:02x}{g:02x}{b:02x}"
        except:
            return color
    
    def toggle_fullscreen(self, event=None):
        """Basculer entre plein écran et fenêtre normale"""
        try:
            if self.root.winfo_viewable():
                current_state = self.root.state()
                if current_state == 'zoomed':
                    self.root.state('normal')
                    self.root.geometry("1400x900")
                    # Centrer la fenêtre
                    self.root.update_idletasks()
                    width = 1400
                    height = 900
                    x = (self.root.winfo_screenwidth() // 2) - (width // 2)
                    y = (self.root.winfo_screenheight() // 2) - (height // 2)
                    self.root.geometry(f"{width}x{height}+{x}+{y}")
                else:
                    self.root.state('zoomed')
        except Exception as e:
            print(f"Erreur toggle fullscreen: {e}")
    
    def toggle_theme(self):
        """Basculer entre le thème sombre et clair"""
        current_mode = ctk.get_appearance_mode()
        if current_mode == "Dark":
            ctk.set_appearance_mode("light")
            self.theme_button.configure(text="☀️ Clair")
        else:
            ctk.set_appearance_mode("dark")
            self.theme_button.configure(text="🌙 Sombre")
    
    def refresh_document_list(self):
        """Rafraîchir la liste des documents avec un style moderne"""
        # Supprimer tous les widgets existants
        for widget in self.document_widgets:
            widget.destroy()
        self.document_widgets.clear()
        
        # Ajouter les documents
        for i, doc in enumerate(self.documents):
            self.create_document_widget(doc, i)
    
    def create_document_widget(self, doc, index):
        """Créer un widget moderne pour un document"""
        # Frame pour le document
        doc_frame = ctk.CTkFrame(self.docs_scrollable, height=80)
        doc_frame.pack(fill="x", padx=5, pady=5)
        doc_frame.pack_propagate(False)
        
        # Frame pour les informations
        info_frame = ctk.CTkFrame(doc_frame, fg_color="transparent")
        info_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Icône selon le type
        type_icons = {
            "PDF": "📄",
            "Word": "📝",
            "Image": "🖼️"
        }
        
        icon = type_icons.get(doc["type"], "📁")
        
        # Nom du document avec icône
        name_label = ctk.CTkLabel(
            info_frame,
            text=f"{icon} {doc['name']}",
            font=self.fonts["body"],
            anchor="w"
        )
        name_label.pack(fill="x", pady=(0, 5))
        
        # Informations détaillées
        details_text = f"Type: {doc['type']} • Taille: {doc['size']} • Pages: {doc['pages']}"
        details_label = ctk.CTkLabel(
            info_frame,
            text=details_text,
            font=self.fonts["caption"],
            text_color=self.colors["text_secondary"],
            anchor="w"
        )
        details_label.pack(fill="x")
        
        # Bouton de sélection
        select_btn = ctk.CTkButton(
            doc_frame,
            text="👁️",
            width=40,
            height=40,
            command=lambda: self.preview_document(doc),
            fg_color=self.colors["accent"],
            hover_color=self.adjust_color(self.colors["accent"], -20)
        )
        select_btn.pack(side="right", padx=10, pady=10)
        
        # Stocker le widget
        self.document_widgets.append(doc_frame)
        
        # Bind pour la sélection
        doc_frame.bind("<Button-1>", lambda e: self.preview_document(doc))
        name_label.bind("<Button-1>", lambda e: self.preview_document(doc))
        details_label.bind("<Button-1>", lambda e: self.preview_document(doc))
    
    def add_pdf(self):
        """Ajouter des fichiers PDF"""
        files = filedialog.askopenfilenames(
            title="Sélectionner des fichiers PDF",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        
        for file_path in files:
            if file_path:
                self.add_document(file_path, "PDF")
    
    def add_word(self):
        """Ajouter des fichiers Word"""
        files = filedialog.askopenfilenames(
            title="Sélectionner des fichiers Word",
            filetypes=[("Word files", "*.docx *.doc"), ("All files", "*.*")]
        )
        
        for file_path in files:
            if file_path:
                self.add_document(file_path, "Word")
    
    def add_images(self):
        """Ajouter des fichiers images"""
        files = filedialog.askopenfilenames(
            title="Sélectionner des images",
            filetypes=[
                ("Image files", "*.png *.jpg *.jpeg *.bmp *.gif *.tiff"),
                ("All files", "*.*")
            ]
        )
        
        for file_path in files:
            if file_path:
                self.add_document(file_path, "Image")
    
    def add_document(self, file_path, doc_type):
        """Ajouter un document à la liste"""
        try:
            file_name = os.path.basename(file_path)
            file_size = self.format_size(os.path.getsize(file_path))
            
            # Compter les pages selon le type
            pages = self.count_pages(file_path, doc_type)
            
            # Ajouter à la liste
            document = {
                "path": file_path,
                "name": file_name,
                "type": doc_type,
                "size": file_size,
                "pages": pages
            }
            
            self.documents.append(document)
            
            # Rafraîchir la liste moderne
            self.refresh_document_list()
            
            self.status_var.set(f"✅ Document ajouté: {file_name}")
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'ajout du document: {str(e)}")
            self.status_var.set(f"❌ Erreur: {str(e)}")
    
    def count_pages(self, file_path, doc_type):
        """Compter le nombre de pages d'un document"""
        try:
            if doc_type == "PDF":
                doc = fitz.open(file_path)
                return len(doc)
            elif doc_type == "Word":
                # Pour Word, on estime (difficile de compter exactement)
                return "~"
            elif doc_type == "Image":
                return 1
            else:
                return "?"
        except:
            return "?"
    
    def format_size(self, size_bytes):
        """Formater la taille du fichier"""
        if size_bytes == 0:
            return "0B"
        
        size_names = ["B", "KB", "MB", "GB"]
        i = 0
        while size_bytes >= 1024 and i < len(size_names) - 1:
            size_bytes /= 1024
            i += 1
        
        return f"{size_bytes:.1f}{size_names[i]}"
    
    def on_document_select(self, event):
        """Gérer la sélection d'un document (méthode obsolète, remplacée par les boutons)"""
        pass
    
    def preview_document(self, document):
        """Afficher l'aperçu d'un document"""
        try:
            self.preview_canvas.delete("all")
            
            if document["type"] == "PDF":
                self.preview_pdf(document["path"])
            elif document["type"] == "Image":
                self.preview_image(document["path"])
            elif document["type"] == "Word":
                self.preview_word(document["path"])
            
        except Exception as e:
            self.preview_canvas.create_text(100, 100, text=f"Erreur d'aperçu: {str(e)}")
    
    def preview_pdf(self, file_path):
        """Aperçu d'un PDF"""
        try:
            doc = fitz.open(file_path)
            page = doc.load_page(0)  # Première page
            
            # Convertir en image
            mat = fitz.Matrix(0.5, 0.5)  # Zoom 50%
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("ppm")
            
            # Afficher dans le canvas
            image = tk.PhotoImage(data=img_data)
            self.preview_canvas.create_image(10, 10, anchor=tk.NW, image=image)
            self.preview_canvas.image = image  # Garder une référence
            
            # Configurer la zone de scroll
            self.preview_canvas.configure(scrollregion=self.preview_canvas.bbox("all"))
            
        except Exception as e:
            self.preview_canvas.create_text(100, 100, text=f"Erreur PDF: {str(e)}")
    
    def preview_image(self, file_path):
        """Aperçu d'une image"""
        try:
            image = Image.open(file_path)
            
            # Redimensionner si nécessaire
            max_size = (400, 400)
            image.thumbnail(max_size, Image.Resampling.LANCZOS)
            
            # Convertir pour tkinter
            photo = ImageTk.PhotoImage(image)
            self.preview_canvas.create_image(10, 10, anchor=tk.NW, image=photo)
            self.preview_canvas.image = photo  # Garder une référence
            
            # Configurer la zone de scroll
            self.preview_canvas.configure(scrollregion=self.preview_canvas.bbox("all"))
            
        except Exception as e:
            self.preview_canvas.create_text(100, 100, text=f"Erreur image: {str(e)}")
    
    def preview_word(self, file_path):
        """Aperçu d'un document Word"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extraire le texte des premiers paragraphes
            for i, paragraph in enumerate(doc.paragraphs[:10]):
                if paragraph.text.strip():
                    text += paragraph.text + "\n\n"
            
            if text:
                self.preview_canvas.create_text(10, 10, anchor=tk.NW, text=text[:500] + "...",
                                              width=380, font=("Arial", 10))
            else:
                self.preview_canvas.create_text(100, 100, text="Document Word vide ou sans texte")
            
            # Configurer la zone de scroll
            self.preview_canvas.configure(scrollregion=self.preview_canvas.bbox("all"))
            
        except Exception as e:
            self.preview_canvas.create_text(100, 100, text=f"Erreur Word: {str(e)}")
    
    def move_up(self):
        """Déplacer le document sélectionné vers le haut"""
        # Trouver le document sélectionné (simplifié pour l'interface moderne)
        if not self.documents:
            messagebox.showwarning("Attention", "Aucun document à déplacer")
            return
        
        # Pour l'instant, on déplace le premier document trouvé
        # TODO: Implémenter une sélection plus sophistiquée
        if len(self.documents) > 1:
            # Déplacer le dernier document vers le haut
            self.documents[0], self.documents[1] = self.documents[1], self.documents[0]
            self.refresh_document_list()
            self.status_var.set("📄 Document déplacé vers le haut")
    
    def move_down(self):
        """Déplacer le document sélectionné vers le bas"""
        # Trouver le document sélectionné (simplifié pour l'interface moderne)
        if not self.documents:
            messagebox.showwarning("Attention", "Aucun document à déplacer")
            return
        
        # Pour l'instant, on déplace le premier document trouvé
        # TODO: Implémenter une sélection plus sophistiquée
        if len(self.documents) > 1:
            # Déplacer le premier document vers le bas
            self.documents[0], self.documents[1] = self.documents[1], self.documents[0]
            self.refresh_document_list()
            self.status_var.set("📄 Document déplacé vers le bas")
    
    def remove_document(self):
        """Supprimer le document sélectionné"""
        if not self.documents:
            messagebox.showwarning("Attention", "Aucun document à supprimer")
            return
        
        if messagebox.askyesno("Confirmation", "Êtes-vous sûr de vouloir supprimer le dernier document ajouté?"):
            # Pour l'instant, on supprime le dernier document
            # TODO: Implémenter une sélection plus sophistiquée
            removed_doc = self.documents.pop()
            self.refresh_document_list()
            self.preview_canvas.delete("all")
            self.status_var.set(f"🗑️ Document supprimé: {removed_doc['name']}")
    
    def refresh_document_list(self):
        """Rafraîchir la liste des documents avec un style moderne"""
        # Supprimer tous les widgets existants
        for widget in self.document_widgets:
            widget.destroy()
        self.document_widgets.clear()
        
        # Ajouter les documents
        for i, doc in enumerate(self.documents):
            self.create_document_widget(doc, i)
    
    def organize_chapters(self):
        """Organiser les documents en chapitres avec une interface moderne"""
        if not self.documents:
            messagebox.showwarning("Attention", "Ajoutez d'abord des documents")
            return
        
        # Créer une fenêtre pour organiser les chapitres
        chapter_window = ctk.CTkToplevel(self.root)
        chapter_window.title("📋 Organiser les chapitres")
        chapter_window.geometry("900x700")
        chapter_window.transient(self.root)
        chapter_window.grab_set()
        
        # Interface de gestion des chapitres
        self.create_chapter_interface(chapter_window)
    
    def create_chapter_interface(self, window):
        """Créer l'interface moderne de gestion des chapitres"""
        main_frame = ctk.CTkFrame(window)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Titre
        title_label = ctk.CTkLabel(
            main_frame,
            text="📋 Organisation des chapitres",
            font=self.fonts["heading"]
        )
        title_label.pack(pady=(20, 10))
        
        # Instructions
        instructions = ctk.CTkLabel(
            main_frame,
            text="Organisez vos documents en chapitres pour une meilleure structure.",
            font=self.fonts["body"],
            text_color=self.colors["text_secondary"]
        )
        instructions.pack(pady=(0, 20))
        
        # Frame pour les boutons
        button_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        button_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        buttons = [
            ("➕ Nouveau chapitre", self.add_chapter, self.colors["success"]),
            ("🗑️ Supprimer chapitre", self.remove_chapter, self.colors["error"]),
            ("✏️ Renommer chapitre", self.rename_chapter, self.colors["warning"])
        ]
        
        for text, command, color in buttons:
            btn = ctk.CTkButton(
                button_frame,
                text=text,
                command=command,
                width=150,
                height=35,
                font=self.fonts["body"],
                fg_color=color,
                hover_color=self.adjust_color(color, -20)
            )
            btn.pack(side="left", padx=5)
        
        # Zone de chapitres (simulation avec un textbox pour l'instant)
        self.chapter_display = ctk.CTkTextbox(
            main_frame,
            height=300,
            font=self.fonts["body"]
        )
        self.chapter_display.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Initialiser avec les documents comme un seul chapitre
        if not self.chapters:
            self.chapters = [{"name": "Chapitre 1", "documents": list(range(len(self.documents)))}]
        
        self.refresh_chapter_display()
        
        # Boutons de validation
        validate_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        validate_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        cancel_btn = ctk.CTkButton(
            validate_frame,
            text="❌ Annuler",
            command=lambda: (self.reset_chapters(), window.destroy()),
            width=100,
            height=35,
            font=self.fonts["body"],
            fg_color=self.colors["error"]
        )
        cancel_btn.pack(side="right", padx=(10, 0))
        
        ok_btn = ctk.CTkButton(
            validate_frame,
            text="✅ OK",
            command=window.destroy,
            width=100,
            height=35,
            font=self.fonts["body"],
            fg_color=self.colors["success"]
        )
        ok_btn.pack(side="right")
    
    def refresh_chapter_display(self):
        """Rafraîchir l'affichage des chapitres"""
        self.chapter_display.delete("0.0", "end")
        
        display_text = "📚 ORGANISATION DES CHAPITRES\n\n"
        
        for i, chapter in enumerate(self.chapters):
            display_text += f"📖 {chapter['name']}\n"
            for doc_index in chapter['documents']:
                if doc_index < len(self.documents):
                    doc = self.documents[doc_index]
                    type_icon = {"PDF": "📄", "Word": "📝", "Image": "🖼️"}.get(doc['type'], "📁")
                    display_text += f"   {type_icon} {doc['name']} ({doc['pages']} pages)\n"
            display_text += "\n"
        
        self.chapter_display.insert("0.0", display_text)
        self.chapter_display.configure(state="disabled")
    
    def add_chapter(self):
        """Ajouter un nouveau chapitre"""
        name = simpledialog.askstring("Nouveau chapitre", "Nom du chapitre:")
        if name:
            self.chapters.append({"name": name, "documents": []})
            self.refresh_chapter_display()
            self.status_var.set(f"➕ Chapitre ajouté: {name}")
    
    def remove_chapter(self):
        """Supprimer un chapitre"""
        if not self.chapters:
            messagebox.showwarning("Attention", "Aucun chapitre à supprimer")
            return
        
        # Créer une liste des chapitres pour sélection
        chapter_names = [chapter["name"] for chapter in self.chapters]
        
        # Demander quel chapitre supprimer
        choice = simpledialog.askstring(
            "Supprimer chapitre", 
            f"Chapitres disponibles: {', '.join(chapter_names)}\nNom du chapitre à supprimer:"
        )
        
        if choice:
            for i, chapter in enumerate(self.chapters):
                if chapter["name"] == choice:
                    if messagebox.askyesno("Confirmation", f"Supprimer le chapitre '{choice}'?"):
                        del self.chapters[i]
                        self.refresh_chapter_display()
                        self.status_var.set(f"🗑️ Chapitre supprimé: {choice}")
                        return
            
            messagebox.showwarning("Attention", "Chapitre non trouvé")
    
    def rename_chapter(self):
        """Renommer un chapitre"""
        if not self.chapters:
            messagebox.showwarning("Attention", "Aucun chapitre à renommer")
            return
        
        # Créer une liste des chapitres pour sélection
        chapter_names = [chapter["name"] for chapter in self.chapters]
        
        # Demander quel chapitre renommer
        old_name = simpledialog.askstring(
            "Renommer chapitre", 
            f"Chapitres disponibles: {', '.join(chapter_names)}\nNom du chapitre à renommer:"
        )
        
        if old_name:
            for chapter in self.chapters:
                if chapter["name"] == old_name:
                    new_name = simpledialog.askstring("Nouveau nom", "Nouveau nom:", initialvalue=old_name)
                    if new_name:
                        chapter["name"] = new_name
                        self.refresh_chapter_display()
                        self.status_var.set(f"✏️ Chapitre renommé: {old_name} → {new_name}")
                        return
            
            messagebox.showwarning("Attention", "Chapitre non trouvé")
    
    def reset_chapters(self):
        """Réinitialiser les chapitres"""
        self.chapters = []
    
    def generate_pdf(self):
        """Générer le PDF final avec une interface moderne"""
        if not self.documents:
            messagebox.showwarning("Attention", "Ajoutez d'abord des documents")
            return
        
        # Choisir le fichier de sortie
        output_file = filedialog.asksaveasfilename(
            title="Sauvegarder le PDF",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            initialdir=self.config["default_output_dir"]
        )
        
        if not output_file:
            return
        
        # Créer une fenêtre de progression moderne
        progress_window = ctk.CTkToplevel(self.root)
        progress_window.title("🔨 Génération du PDF")
        progress_window.geometry("500x200")
        progress_window.transient(self.root)
        progress_window.grab_set()
        
        # Frame principal
        progress_frame = ctk.CTkFrame(progress_window)
        progress_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Titre
        title_label = ctk.CTkLabel(
            progress_frame,
            text="🔨 Génération du PDF en cours...",
            font=self.fonts["heading"]
        )
        title_label.pack(pady=(20, 10))
        
        # Label de progression
        progress_label = ctk.CTkLabel(
            progress_frame,
            text="Préparation...",
            font=self.fonts["body"]
        )
        progress_label.pack(pady=(0, 20))
        
        # Barre de progression
        progress_bar = ctk.CTkProgressBar(
            progress_frame,
            width=400,
            height=20
        )
        progress_bar.pack(pady=(0, 20))
        progress_bar.set(0)
        
        # Bouton d'annulation
        cancel_button = ctk.CTkButton(
            progress_frame,
            text="❌ Annuler",
            width=100,
            height=35,
            font=self.fonts["body"],
            fg_color=self.colors["error"]
        )
        cancel_button.pack(pady=(0, 20))
        
        # Lancer la génération dans un thread
        self.generation_cancelled = False
        
        def on_cancel():
            self.generation_cancelled = True
            progress_window.destroy()
        
        cancel_button.configure(command=on_cancel)
        
        def generate_thread():
            try:
                self.generate_pdf_content(output_file, progress_label, progress_bar)
                if not self.generation_cancelled:
                    progress_window.after(0, lambda: (
                        progress_window.destroy(),
                        messagebox.showinfo("Succès", f"🎉 PDF généré avec succès!\n📁 {output_file}")
                    ))
            except Exception as e:
                if not self.generation_cancelled:
                    error_msg = str(e)  # Capturer la variable avant la lambda
                    progress_window.after(0, lambda: (
                        progress_window.destroy(),
                        messagebox.showerror("Erreur", f"❌ Erreur lors de la génération: {error_msg}")
                    ))
        
        thread = threading.Thread(target=generate_thread)
        thread.daemon = True
        thread.start()
    
    def generate_pdf_content(self, output_file, progress_label, progress_bar):
        """Générer le contenu du PDF avec feedback moderne"""
        # Préparer les styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Centre
            textColor=colors.darkblue
        )
        chapter_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=20,
            spaceBefore=20,
            textColor=colors.darkred
        )
        # Créer le document PDF
        pdf_doc = SimpleDocTemplate(output_file, pagesize=A4)
        story = []
        temp_images = []  # Stocker les fichiers temporaires à supprimer après
        
        # Page de titre
        progress_label.configure(text="📄 Création de la page de titre...")
        progress_bar.set(0.1)
        
        title = Paragraph("📚 Document Relié", title_style)
        story.append(title)
        story.append(Spacer(1, 0.5*inch))
        
        # Date de création
        date_text = f"📅 Créé le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        story.append(Paragraph(date_text, styles['Normal']))
        story.append(Spacer(1, 0.5*inch))
        
        # Table des matières
        progress_label.configure(text="📋 Génération de la table des matières...")
        progress_bar.set(0.2)
        
        toc_title = Paragraph("📋 Table des matières", chapter_style)
        story.append(toc_title)
        
        # Si des chapitres sont définis, les utiliser
        if self.chapters:
            for i, chapter in enumerate(self.chapters):
                chapter_text = f"{i+1}. {chapter['name']}"
                story.append(Paragraph(chapter_text, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        else:
            # Sinon, lister tous les documents
            for i, document in enumerate(self.documents):
                doc_text = f"{i+1}. {document['name']}"
                story.append(Paragraph(doc_text, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        story.append(PageBreak())
        
        # Contenu des documents
        total_docs = len(self.documents)
        
        if self.chapters:
            # Traiter par chapitres
            for chapter_num, chapter in enumerate(self.chapters):
                if self.generation_cancelled:
                    return
                
                progress_label.configure(text=f"📖 Traitement du chapitre: {chapter['name']}")
                progress_bar.set(0.3 + (chapter_num / len(self.chapters)) * 0.6)
                
                # Titre du chapitre
                chapter_title = Paragraph(f"📖 Chapitre {chapter_num + 1}: {chapter['name']}", chapter_style)
                story.append(chapter_title)
                story.append(Spacer(1, 0.3*inch))
                
                # Documents du chapitre
                for doc_index in chapter['documents']:
                    if doc_index < len(self.documents):
                        document = self.documents[doc_index]
                        self.add_document_to_story(story, document, styles, temp_images)
                
                story.append(PageBreak())
        else:
            # Traiter tous les documents
            for i, document in enumerate(self.documents):
                if self.generation_cancelled:
                    return
                
                progress_label.configure(text=f"📄 Traitement de: {document['name']}")
                progress_bar.set(0.3 + (i / total_docs) * 0.6)
                
                self.add_document_to_story(story, document, styles, temp_images)
        
        # Construire le PDF
        progress_label.configure(text="🔨 Finalisation du PDF...")
        progress_bar.set(0.95)
        
        pdf_doc.build(story)
        
        # Supprimer les fichiers temporaires après la génération
        for temp_path in temp_images:
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except Exception:
                pass
        
        progress_bar.set(1.0)
        progress_label.configure(text="✅ PDF généré avec succès!")
    
    def add_document_to_story(self, story, doc, styles, temp_images=None):
        """Ajouter un document à l'histoire du PDF"""
        try:
            # Titre du document
            doc_title = Paragraph(f"Document: {doc['name']}", styles['Heading3'])
            story.append(doc_title)
            story.append(Spacer(1, 0.2*inch))
            
            if doc['type'] == 'PDF':
                self.add_pdf_to_story(story, doc['path'], temp_images)
            elif doc['type'] == 'Word':
                self.add_word_to_story(story, doc['path'], styles)
            elif doc['type'] == 'Image':
                self.add_image_to_story(story, doc['path'])
            
            story.append(Spacer(1, 0.3*inch))
            
        except Exception as e:
            # En cas d'erreur, ajouter un message d'erreur
            error_text = f"Erreur lors du traitement de {doc['name']}: {str(e)}"
            story.append(Paragraph(error_text, styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
    
    def add_pdf_to_story(self, story, pdf_path, temp_images=None):
        """Ajouter un PDF à l'histoire"""
        if temp_images is None:
            temp_images = []
        try:
            pdf_doc = fitz.open(pdf_path)
            
            # Créer un identifiant unique pour ce PDF basé sur son nom
            import time
            pdf_id = f"{os.path.basename(pdf_path).replace('.', '_')}_{int(time.time() * 1000)}"
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.load_page(page_num)
                
                # Convertir en image
                mat = fitz.Matrix(1.5, 1.5)  # Zoom 150%
                pix = page.get_pixmap(matrix=mat)
                
                # Sauvegarder temporairement avec un nom unique
                temp_path = f"temp_{pdf_id}_page_{page_num}.png"
                pix.save(temp_path)
                temp_images.append(temp_path)
                
                # Ajouter à l'histoire
                img = RLImage(temp_path, width=6*inch, height=8*inch)
                story.append(img)
                
                if page_num < len(pdf_doc) - 1:
                    story.append(PageBreak())
            
            pdf_doc.close()
            
        except Exception as e:
            # Obtenir les styles depuis getSampleStyleSheet
            styles = getSampleStyleSheet()
            story.append(Paragraph(f"Erreur PDF: {str(e)}", styles['Normal']))
    
    def add_word_to_story(self, story, word_path, styles):
        """Ajouter un document Word à l'histoire"""
        try:
            doc = Document(word_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Déterminer le style en fonction du formatage
                    if paragraph.style.name.startswith('Heading'):
                        style = styles['Heading2']
                    else:
                        style = styles['Normal']
                    
                    p = Paragraph(paragraph.text, style)
                    story.append(p)
                    story.append(Spacer(1, 0.1*inch))
            
        except Exception as e:
            story.append(Paragraph(f"Erreur Word: {str(e)}", styles['Normal']))
    
    def add_image_to_story(self, story, image_path):
        """Ajouter une image à l'histoire"""
        try:
            # Ouvrir l'image pour obtenir ses dimensions
            with Image.open(image_path) as img:
                width, height = img.size
                
                # Calculer les dimensions pour s'adapter à la page
                max_width = 6*inch
                max_height = 8*inch
                
                ratio = min(max_width/width, max_height/height, 1)
                new_width = width * ratio
                new_height = height * ratio
                
                # Ajouter l'image
                rl_img = RLImage(image_path, width=new_width, height=new_height)
                story.append(rl_img)
                
        except Exception as e:
            # Obtenir les styles depuis getSampleStyleSheet
            styles = getSampleStyleSheet()
            story.append(Paragraph(f"Erreur image: {str(e)}", styles['Normal']))
    
    def new_project(self):
        """Nouveau projet"""
        if self.documents:
            if messagebox.askyesno("Nouveau projet", "Voulez-vous vraiment créer un nouveau projet?\nLes modifications non sauvegardées seront perdues."):
                self.documents = []
                self.chapters = []
                self.refresh_document_list()
                self.preview_canvas.delete("all")
                self.status_var.set("Nouveau projet créé")
    
    def save_project(self):
        """Sauvegarder le projet"""
        if not self.documents:
            messagebox.showwarning("Attention", "Aucun document à sauvegarder")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Sauvegarder le projet",
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                project_data = {
                    "documents": self.documents,
                    "chapters": self.chapters,
                    "created": datetime.now().isoformat(),
                    "version": "1.0"
                }
                
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(project_data, f, ensure_ascii=False, indent=2)
                
                messagebox.showinfo("Succès", "Projet sauvegardé avec succès!")
                self.status_var.set(f"Projet sauvegardé: {os.path.basename(file_path)}")
                
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de la sauvegarde: {str(e)}")
    
    def open_project(self):
        """Ouvrir un projet"""
        file_path = filedialog.askopenfilename(
            title="Ouvrir un projet",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    project_data = json.load(f)
                
                # Vérifier que les fichiers existent toujours
                valid_documents = []
                for doc in project_data.get("documents", []):
                    if os.path.exists(doc["path"]):
                        valid_documents.append(doc)
                    else:
                        messagebox.showwarning("Fichier manquant", f"Le fichier {doc['name']} est introuvable")
                
                self.documents = valid_documents
                self.chapters = project_data.get("chapters", [])
                
                self.refresh_document_list()
                self.preview_canvas.delete("all")
                
                messagebox.showinfo("Succès", "Projet ouvert avec succès!")
                self.status_var.set(f"Projet ouvert: {os.path.basename(file_path)}")
                
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de l'ouverture: {str(e)}")
    
    def show_settings(self):
        """Afficher les paramètres avec une interface moderne"""
        settings_window = ctk.CTkToplevel(self.root)
        settings_window.title("⚙️ Paramètres")
        settings_window.geometry("500x400")
        settings_window.transient(self.root)
        settings_window.grab_set()
        
        # Frame principal
        main_frame = ctk.CTkFrame(settings_window)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Titre
        title_label = ctk.CTkLabel(
            main_frame,
            text="⚙️ Paramètres de l'application",
            font=self.fonts["heading"]
        )
        title_label.pack(pady=(20, 30))
        
        # Dossier de sortie par défaut
        output_label = ctk.CTkLabel(
            main_frame,
            text="📁 Dossier de sortie par défaut:",
            font=self.fonts["body"],
            anchor="w"
        )
        output_label.pack(fill="x", padx=20, pady=(0, 10))
        
        output_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        output_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        self.output_dir_var = tk.StringVar(value=self.config["default_output_dir"])
        output_entry = ctk.CTkEntry(
            output_frame,
            textvariable=self.output_dir_var,
            height=35,
            font=self.fonts["body"]
        )
        output_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))
        
        browse_btn = ctk.CTkButton(
            output_frame,
            text="🔍 Parcourir",
            command=self.browse_output_dir,
            width=100,
            height=35,
            font=self.fonts["body"]
        )
        browse_btn.pack(side="right")
        
        # Taille de page
        page_label = ctk.CTkLabel(
            main_frame,
            text="📄 Taille de page:",
            font=self.fonts["body"],
            anchor="w"
        )
        page_label.pack(fill="x", padx=20, pady=(0, 10))
        
        self.page_size_var = tk.StringVar(value=self.config["page_size"])
        page_combo = ctk.CTkComboBox(
            main_frame,
            values=["A4", "Letter", "Legal"],
            variable=self.page_size_var,
            height=35,
            font=self.fonts["body"]
        )
        page_combo.pack(fill="x", padx=20, pady=(0, 30))
        
        # Boutons
        button_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        button_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        cancel_btn = ctk.CTkButton(
            button_frame,
            text="❌ Annuler",
            command=settings_window.destroy,
            width=100,
            height=35,
            font=self.fonts["body"],
            fg_color=self.colors["error"]
        )
        cancel_btn.pack(side="right", padx=(10, 0))
        
        ok_btn = ctk.CTkButton(
            button_frame,
            text="✅ OK",
            command=lambda: self.save_settings(settings_window),
            width=100,
            height=35,
            font=self.fonts["body"],
            fg_color=self.colors["success"]
        )
        ok_btn.pack(side="right")
    
    def browse_output_dir(self):
        """Parcourir pour choisir le dossier de sortie"""
        directory = filedialog.askdirectory(initialdir=self.output_dir_var.get())
        if directory:
            self.output_dir_var.set(directory)
    
    def save_settings(self, window):
        """Sauvegarder les paramètres"""
        self.config["default_output_dir"] = self.output_dir_var.get()
        self.config["page_size"] = self.page_size_var.get()
        
        # Créer le dossier s'il n'existe pas
        os.makedirs(self.config["default_output_dir"], exist_ok=True)
        
        window.destroy()
        messagebox.showinfo("Succès", "Paramètres sauvegardés!")
    
    def show_help(self):
        """Afficher l'aide avec une interface moderne"""
        help_text = """
🚀 GUIDE D'UTILISATION - RELIEUR VIRTUEL DE DOCUMENTS

📄 1. AJOUTER DES DOCUMENTS
   • Utilisez les boutons "Ajouter PDF", "Ajouter Word", "Ajouter Images"
   • Sélectionnez un ou plusieurs fichiers à la fois
   • Les documents apparaissent dans la liste de gauche

📋 2. ORGANISER LES DOCUMENTS
   • Utilisez "Monter" et "Descendre" pour changer l'ordre
   • Cliquez sur "Organiser Chapitres" pour créer des chapitres
   • Cliquez sur l'icône 👁️ pour prévisualiser un document

👁️ 3. PRÉVISUALISER
   • Cliquez sur un document dans la liste pour le prévisualiser
   • L'aperçu apparaît dans le panneau de droite

🔨 4. GÉNÉRER LE PDF
   • Cliquez sur "Générer PDF"
   • Choisissez l'emplacement de sauvegarde
   • Le PDF final inclut une table des matières

💾 5. GÉRER LES PROJETS
   • Menu Fichier > Nouveau projet
   • Menu Fichier > Sauvegarder projet
   • Menu Fichier > Ouvrir projet

📁 FORMATS SUPPORTÉS:
   • PDF: .pdf
   • Word: .docx, .doc
   • Images: .png, .jpg, .jpeg, .bmp, .gif, .tiff

💡 CONSEILS:
   • Organisez vos documents avant de générer le PDF
   • Utilisez des noms de chapitres descriptifs
   • Vérifiez l'aperçu avant la génération finale
   • Basculez entre les thèmes sombre et clair avec le bouton en haut à droite
        """
        
        help_window = ctk.CTkToplevel(self.root)
        help_window.title("📚 Guide d'utilisation")
        help_window.geometry("700x600")
        help_window.transient(self.root)
        
        # Frame principal
        main_frame = ctk.CTkFrame(help_window)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Titre
        title_label = ctk.CTkLabel(
            main_frame,
            text="📚 Guide d'utilisation",
            font=self.fonts["heading"]
        )
        title_label.pack(pady=(20, 20))
        
        # Texte d'aide
        text_widget = ctk.CTkTextbox(
            main_frame,
            font=self.fonts["body"],
            wrap="word"
        )
        text_widget.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Insérer le texte
        text_widget.insert("0.0", help_text)
        text_widget.configure(state="disabled")
        
        # Bouton fermer
        close_btn = ctk.CTkButton(
            main_frame,
            text="✅ Fermer",
            command=help_window.destroy,
            width=100,
            height=35,
            font=self.fonts["body"],
            fg_color=self.colors["primary"]
        )
        close_btn.pack(pady=(0, 20))
    
    def show_about(self):
        """Afficher les informations sur l'application"""
        about_text = """
RELIEUR VIRTUEL DE DOCUMENTS
Version 1.0

Développé pour assembler facilement plusieurs documents
en un seul PDF professionnel avec table des matières.

Fonctionnalités:
• Support PDF, Word et Images
• Organisation en chapitres
• Table des matières automatique
• Aperçu des documents
• Sauvegarde de projets

© 2025 - Relieur Virtuel
        """
        
        messagebox.showinfo("À propos", about_text)

def main():
    """Fonction principale avec interface moderne"""
    try:
        # Initialiser CustomTkinter
        root = ctk.CTk()
        app = DocumentBinder(root)
        
        # Centrer la fenêtre
        root.update_idletasks()
        width = root.winfo_width()
        height = root.winfo_height()
        x = (root.winfo_screenwidth() // 2) - (width // 2)
        y = (root.winfo_screenheight() // 2) - (height // 2)
        root.geometry(f"{width}x{height}+{x}+{y}")
        
        # Lancer l'application
        root.mainloop()
        
    except Exception as e:
        messagebox.showerror("Erreur", f"Erreur lors du démarrage: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()